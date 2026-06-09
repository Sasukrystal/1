# -*- coding: utf-8 -*-
"""Generate CG course report with embedded screenshots and captioned tables/figures."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import shutil

REPORT_ROOT = r"C:\Users\刘昌松\Desktop\黑暗地牢"
OUTPUT = os.path.join(REPORT_ROOT, "计算机图形学大作业报告_黑暗地牢.docx")
SCREENSHOT_DIR = os.path.join(REPORT_ROOT, "screenshots")
MANUAL_SCREENSHOT_DIR = os.path.join(SCREENSHOT_DIR, "手动截图")
FIG7_3_COMPOSITE = os.path.join(SCREENSHOT_DIR, "fig7-3_composite.png")
FIG7_4_COMPOSITE = os.path.join(SCREENSHOT_DIR, "fig7-4_composite.png")
SCREENSHOT_SOURCE = r"d:\unity\test\bagsys\Assets\Screenshots"

_table_idx = {}
_figure_idx = {}


def set_cn_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_cn_font(run, "黑体", sizes.get(level, 12), bold=True)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_cn_font(run, "宋体", 12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, "宋体", 12)
    return p


def next_table_no(chapter):
    _table_idx[chapter] = _table_idx.get(chapter, 0) + 1
    return _table_idx[chapter]


def next_figure_no(chapter):
    _figure_idx[chapter] = _figure_idx.get(chapter, 0) + 1
    return _figure_idx[chapter]


def add_table_caption(doc, chapter, title):
    no = next_table_no(chapter)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"表 {chapter}-{no}  {title}")
    set_cn_font(run, "黑体", 10.5, bold=True)


def add_figure_caption(doc, chapter, no, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"图 {chapter}-{no}  {title}")
    set_cn_font(run, "黑体", 10.5, bold=True)


def _append_cell_border(tc_borders, edge, spec):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), spec["val"])
    if spec["val"] != "nil":
        el.set(qn("w:sz"), spec["sz"])
        el.set(qn("w:space"), spec["space"])
        el.set(qn("w:color"), spec["color"])
    tc_borders.append(el)


def apply_three_line_table(table):
    """三线表：顶线、表头下线、底线；无竖线与中间横线。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    old_tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if old_tbl_borders is not None:
        tbl_pr.remove(old_tbl_borders)

    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    top_line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_line = {"val": "single", "sz": "6", "space": "0", "color": "000000"}
    bot_line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    none = {"val": "nil"}

    num_rows = len(table.rows)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            old = tc_pr.find(qn("w:tcBorders"))
            if old is not None:
                tc_pr.remove(old)

            tc_borders = OxmlElement("w:tcBorders")
            if num_rows == 1:
                spec = {"top": top_line, "bottom": bot_line, "left": none, "right": none}
            elif ri == 0:
                spec = {"top": top_line, "bottom": mid_line, "left": none, "right": none}
            elif ri == num_rows - 1:
                spec = {"top": none, "bottom": bot_line, "left": none, "right": none}
            else:
                spec = {"top": none, "bottom": none, "left": none, "right": none}

            for edge, border_spec in spec.items():
                _append_cell_border(tc_borders, edge, border_spec)
            tc_pr.append(tc_borders)


def add_table(doc, chapter, caption, headers, rows):
    add_table_caption(doc, chapter, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_cn_font(r, "宋体", 10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    set_cn_font(r, "宋体", 10)
    apply_three_line_table(table)
    doc.add_paragraph()


def resolve_image(filename):
    candidates = [
        os.path.join(SCREENSHOT_DIR, filename),
        os.path.join(MANUAL_SCREENSHOT_DIR, filename),
        os.path.join(SCREENSHOT_SOURCE, filename),
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def build_figure_7_3_composite():
    """将 7-3-1 / 7-3-2 / 7-3-3 拼成图 7-3 插入报告。"""
    from PIL import Image, ImageDraw, ImageFont

    names = ["7-3-1射击.png", "7-3-2boss战.png", "7-3-3弓箭手拉弓.png"]
    paths = [os.path.join(MANUAL_SCREENSHOT_DIR, n) for n in names]
    for p in paths:
        if not os.path.isfile(p):
            print("Warning: missing for 7-3 composite:", p)
            return False

    gap = 12
    label_h = 28
    bg = (248, 248, 248)

    img1 = Image.open(paths[0]).convert("RGB")
    img2 = Image.open(paths[1]).convert("RGB")
    img3 = Image.open(paths[2]).convert("RGB")

    def resize_to_height(img, height):
        width = max(1, int(img.width * height / img.height))
        return img.resize((width, height), Image.LANCZOS)

    top_h = 220
    top1 = resize_to_height(img1, top_h)
    top3 = resize_to_height(img3, top_h)
    top_w = top1.width + gap + top3.width

    bottom_w = max(top_w, 900)
    bottom_h = int(img2.height * bottom_w / img2.width)
    bottom = img2.resize((bottom_w, bottom_h), Image.LANCZOS)

    # 顶行居中于底图宽度
    top_row = Image.new("RGB", (bottom_w, top_h + label_h), bg)
    x1 = (bottom_w - top_w) // 2
    top_row.paste(top1, (x1, label_h))
    top_row.paste(top3, (x1 + top1.width + gap, label_h))

    total_h = label_h + top_h + gap + bottom_h + label_h
    canvas = Image.new("RGB", (bottom_w, total_h), bg)
    canvas.paste(top_row, (0, 0))
    canvas.paste(bottom, (0, label_h + top_h + gap))

    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("msyh.ttc", 18)
    except OSError:
        font = ImageFont.load_default()

    labels = [
        (x1 + top1.width // 2, 4, "(a) 箭矢投射物透明边缘"),
        (x1 + top1.width + gap + top3.width // 2, 4, "(b) 弓箭手蓄力姿态"),
        (bottom_w // 2, label_h + top_h + gap + bottom_h + 4, "(c) Boss 战技能特效与精灵显示"),
    ]
    for cx, y, text in labels:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((cx - tw // 2, y), text, fill=(40, 40, 40), font=font)

    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    canvas.save(FIG7_3_COMPOSITE, quality=95)
    print("Built 7-3 composite:", FIG7_3_COMPOSITE, canvas.size)
    return True


def build_figure_7_4_composite():
    """将 7-4-1~4 四张 UI 截图拼成 2×2 网格作为图 7-4。"""
    from PIL import Image, ImageDraw, ImageFont

    names = [
        "7-4-1背包界面.png",
        "7-4-2角色面板界面.png",
        "7-4-3虫核界面.png",
        "7-4-4宝物界面.png",
    ]
    labels = [
        "(a) 背包界面",
        "(b) 角色面板",
        "(c) 虫核界面",
        "(d) 宝物界面",
    ]
    paths = [os.path.join(MANUAL_SCREENSHOT_DIR, n) for n in names]
    for p in paths:
        if not os.path.isfile(p):
            print("Warning: missing for 7-4 composite:", p)
            return False

    gap = 10
    label_h = 26
    cell_w = 580
    cell_h = int(cell_w * 720 / 1188)
    bg = (248, 248, 248)

    cols, rows = 2, 2
    canvas_w = cols * cell_w + (cols + 1) * gap
    canvas_h = rows * (cell_h + label_h) + (rows + 1) * gap
    canvas = Image.new("RGB", (canvas_w, canvas_h), bg)
    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("msyh.ttc", 17)
    except OSError:
        font = ImageFont.load_default()

    for idx, (path, label) in enumerate(zip(paths, labels)):
        col = idx % cols
        row = idx // cols
        x0 = gap + col * (cell_w + gap)
        y0 = gap + row * (cell_h + label_h + gap)
        img = Image.open(path).convert("RGB")
        img = img.resize((cell_w, cell_h), Image.LANCZOS)
        canvas.paste(img, (x0, y0 + label_h))
        bbox = draw.textbbox((0, 0), label, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x0 + (cell_w - tw) // 2, y0 + 2), label, fill=(40, 40, 40), font=font)

    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    canvas.save(FIG7_4_COMPOSITE, quality=95)
    print("Built 7-4 composite:", FIG7_4_COMPOSITE, canvas.size)
    return True


def add_figure(doc, chapter, title, filename, width_cm=14.0):
    no = next_figure_no(chapter)
    path = resolve_image(filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    if path:
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
    else:
        run = p.add_run(f"[缺少截图：{filename}]")
        set_cn_font(run, "宋体", 10)
    add_figure_caption(doc, chapter, no, title)
    return no


def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    lines = [
        ("学    号", "1023005941"),
        ("题    目", "基于Unity的2D Roguelike地牢游戏《黑暗地牢》设计与实现"),
        ("学    院", "数统学院"),
        ("专    业", "信计"),
        ("班    级", "信计2302"),
        ("姓    名", "刘昌松"),
        ("指导教师", "陆济湘"),
    ]
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}    {value}")
        set_cn_font(r1, "宋体", 14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月26日")
    set_cn_font(r, "宋体", 14)
    doc.add_page_break()


def sync_screenshots_to_report_folder():
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    if not os.path.isdir(SCREENSHOT_SOURCE):
        print("Warning: screenshot source missing:", SCREENSHOT_SOURCE)
        return
    count = 0
    for name in os.listdir(SCREENSHOT_SOURCE):
        if name.lower().endswith(".png"):
            shutil.copy2(os.path.join(SCREENSHOT_SOURCE, name), os.path.join(SCREENSHOT_DIR, name))
            count += 1
    print(f"Synced {count} screenshots to:", SCREENSHOT_DIR)


def build_report():
    global _table_idx, _figure_idx
    _table_idx = {}
    _figure_idx = {}

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 正文从摘要开始（封面由学校统一模板单独提交，不写入本文件）

    add_title(doc, "摘  要", 1)
    add_body(
        doc,
        "本文介绍一款基于 Unity 2022.3 开发的 2D 俯视角 Roguelike 地牢游戏《黑暗地牢》。"
        "项目在图形学层面采用 SpriteRenderer 精灵渲染、序列帧与 Animator 双轨动画、"
        "透明贴图处理及运行时 UI 构建；在系统设计层面实现程序化地牢、五层难度递进、"
        "虫核元素套装与 Meta 成长等机制。",
        indent=True,
    )
    add_body(
        doc,
        "论文重点阐述系统架构、功能模块与类设计、游戏规则与进度系统、"
        "图形渲染与场景动画实现方法，并结合运行截图验证 Editor 与 Release 版本的一致性。"
        "实验表明，游戏已形成从安全大厅、多层地牢探索到 Boss 战与 Windows 独立发布的完整闭环。",
        indent=True,
    )
    add_body(doc, "关键词：计算机图形学；Unity；2D 游戏；Roguelike；精灵动画；程序化地牢", indent=True)
    doc.add_page_break()

    add_title(doc, "第一章  绪论", 1)
    add_title(doc, "1.1  研究目的", 2)
    add_body(
        doc,
        "本课程大作业旨在将计算机图形学理论知识与游戏开发实践相结合，通过完整实现一款可运行的 2D 游戏，"
        "深入理解二维图形渲染、精灵贴图、帧动画、UI 绘制、资源管理与场景组织等核心内容。"
        "具体目标包括：（1）掌握 Unity 引擎中 SpriteRenderer、Physics2D、Canvas 等图形与交互组件的使用；"
        "（2）实现角色、敌人、Boss、环境装饰等视觉元素的加载与显示；"
        "（3）设计并实现一套可扩展的游戏架构，使玩法逻辑与渲染逻辑解耦；"
        "（4）完成 Windows 平台发布，验证 Editor 与 Release 版本在资源加载与渲染上的一致性。",
        indent=True,
    )
    add_title(doc, "1.2  研究意义", 2)
    add_body(
        doc,
        "从图形学教学角度看，2D 游戏是理解光栅化显示、纹理映射、Alpha 混合、排序层（Sorting Layer）"
        "与动画插值等概念的直观载体。与纯理论作业相比，以完整游戏为载体的实践能够暴露真实工程问题，"
        "例如：贴图透明通道处理、非可读纹理在发布包中的切片失败、序列帧 pivot 对齐、"
        "多分辨率 UI 适配等。本项目在开发过程中逐一解决上述问题，具有较好的工程训练价值。"
        "从游戏设计角度看，Roguelike 地牢游戏强调随机性、单局成长与重复可玩性，"
        "适合展示程序化内容生成（PCG）与数据驱动设计思想。",
        indent=True,
    )
    add_title(doc, "1.3  论文结构", 2)
    add_body(
        doc,
        "本文其余章节安排如下：第二章介绍开发环境与相关技术；第三章阐述系统总体设计；"
        "第四章详细说明各功能模块及核心类；第五章说明游戏规则与操作方法；"
        "第六章分析难度、局外/局内双构筑、虫核套装与进度系统；第七章重点论述图形渲染、场景与动画实现；"
        "第八章对比同类游戏并总结创新点；第九章展示运行截图与测试；第十章给出结论与展望。",
        indent=True,
    )

    add_title(doc, "第二章  开发环境与相关技术", 1)
    add_title(doc, "2.1  开发软件与工具", 2)
    add_table(
        doc, "2", "开发软件与工具",
        ["软件/工具", "版本/说明", "用途"],
        [
            ["Unity Editor", "2022.3.57f1c2 LTS", "游戏引擎、场景编辑、构建发布"],
            ["C# / .NET", "Mono2x 脚本后端", "游戏逻辑实现"],
            ["Visual Studio / Cursor", "—", "代码编辑与调试"],
            ["TextMeshPro", "3.0.7", "高质量 UI 文字渲染"],
            ["Physics2D", "Unity 内置", "2D 碰撞与移动"],
            ["Built-in Render Pipeline", "—", "2D 精灵渲染（非 URP/HDRP）"],
            ["Aseprite / 外部像素包", "Tiny Swords 等", "角色与环境美术资源"],
            ["Git", "—", "版本管理"],
        ],
    )
    add_title(doc, "2.2  图形学相关技术", 2)
    add_body(doc, "本项目图形实现以 2D 精灵渲染为主，涉及以下图形学相关技术点：", indent=True)
    add_bullet(doc, "纹理与精灵：PNG 贴图导入 Unity，通过 SpriteRenderer 显示；支持 Point 过滤以保持像素风格。")
    add_bullet(doc, "Alpha 混合与透明抠图：对投射物、Boss 特效等使用边缘采样与泛洪填充去除白底/色键背景。")
    add_bullet(doc, "序列帧动画：将横向 sprite sheet 按固定帧宽切片，LateUpdate 驱动换帧。")
    add_bullet(doc, "Animator 状态机：战士职业使用 Mecanim 控制 Idle/Run/Attack/Guard。")
    add_bullet(doc, "排序与深度：Sorting Layer + sortingOrder 控制玩家、敌人、UI 的绘制顺序。")
    add_bullet(doc, "程序化 UI：运行时动态创建 Canvas、Image、Text，统一暗色地牢风格。")
    add_bullet(doc, "碰撞与视觉分离：物理碰撞壳与装饰性 Sprite 分属不同 GameObject，便于独立迭代美术。")

    add_title(doc, "第三章  系统总体设计", 1)
    add_title(doc, "3.1  架构概述", 2)
    add_body(
        doc,
        "游戏采用「导演（Director）+ 数据模型（Model）+ 视觉（Visual）+ UI 面板（Panel）」的分层架构。"
        "运行时由 ModernRogueBootstrapper 自举各子系统，禁用早期 Bagsys.RogueLike 命名空间下的 3D 地牢与旧 UI，"
        "统一迁移到 ModernRogue 命名空间下的 2D 俯视角实现。核心单例 SoulKnightDirector 负责整局流程："
        "主菜单、大厅、五层地牢切换、HUD、结算与存档。",
        indent=True,
    )
    add_title(doc, "3.2  主循环流程", 2)
    add_body(
        doc,
        "玩家启动游戏 → 主菜单（新游戏/继续/设置/帮助）→ 选择职业 → 进入安全大厅 → "
        "与雕像/武器铺/宝箱/传送门交互 → 进入地牢 → 探索房间、战斗、拾取奖励 → 击败 Boss → "
        "进入下一层或通关 → 死亡或胜利结算。局内数据通过 GameSaveService 写入 PlayerPrefs，支持断点续玩。",
        indent=True,
    )
    add_title(doc, "3.3  目录与资源组织", 2)
    add_table(
        doc, "3", "项目目录与资源组织",
        ["路径", "说明"],
        [
            ["Assets/Scenes/Main.unity", "唯一发布场景"],
            ["Assets/Scripts/ModernRogue/", "主逻辑脚本（77 个 C# 文件）"],
            ["Assets/Resources/Art2D/", "2D 精灵、动画序列、UI、投射物"],
            ["Assets/Resources/ArtIntegration/", "运行时加载的环境/玩家 Prefab"],
            ["Assets/StreamingAssets/", "音频、弓手图集镜像（发布包直读）"],
            ["Assets/Screenshots/", "开发验证截图"],
            ["Builds/Windows/", "Windows 独立发布输出"],
        ],
    )
    add_title(doc, "3.4  程序化地牢生成算法", 2)
    add_body(
        doc,
        "本游戏每一层地牢的拓扑结构并非手工摆放，而是由 SoulKnightDungeonBuilder.GenerateFloorLayout "
        "在运行时按 StageProfile 参数随机生成。算法输入为当前层数 stage 与对应 SoulKnightStageProfile，"
        "输出为一组带网格坐标、房间类型与奖励类型的 GeneratedRoom 列表，再经 GridToWorld 映射为世界坐标并实例化碰撞与视觉。",
        indent=True,
    )
    add_body(
        doc,
        "生成过程可概括为以下步骤："
        "（1）**起点放置**：在二维整数网格 (0,−1) 放置 Start 房；"
        "（2）**随机游走扩展**：从当前节点随机选取上/下/左/右之一扩展新房间，"
        "直至房间数达到 RoomCountMin~Max；边界约束为 y≥−1、|x|≤3、y≤3，防止地图无限扩张；"
        "（3）**Boss 房标记**：计算各房间到起点的曼哈顿距离，将最远节点标记为 Boss 房；"
        "（4）**特殊房分配**：对其余节点按 ShopChance、TreasureChance、CrossroadChance "
        "依次尝试标记商店、宝藏或岔路，未命中者默认为战斗房；"
        "（5）**奖励权重**：战斗房按 StageProfile 中 Gold/Vitality/Core/Treasure/Equipment 权重随机分配 RoomRewardType；"
        "（6）**世界映射**：房间中心间距 RoomGap=36，单房约 30×20，走廊门洞宽度 7.6，"
        "相邻房间通过 LinkGeneratedRooms 建立连通关系并生成走廊碰撞。",
        indent=True,
    )
    add_body(
        doc,
        "图 3-1 为开发调试用的「上帝视角」房间一览：每个矩形块代表一个网格房间，"
        "连线表示可通行的走廊拓扑。可以直观看到本算法产生的分支结构——"
        "同层每次进入地牢时房间数量、相对位置与连接关系均不同，"
        "从而保证 Roguelike 所要求的重玩价值；"
        "同时 StageProfile 又约束了房间数范围与特殊房概率，使难度曲线可控。",
        indent=True,
    )
    add_figure(doc, "3", "程序化地牢网格拓扑（上帝视角房间一览）", "上帝视角房间一览.png", 14)
    add_title(doc, "3.5  伤害与战斗数值", 2)
    add_body(
        doc,
        "基础伤害公式：实际伤害 = max(1, 攻击方攻击力 − (防守方防御力 − 湿润减防))。"
        "战士格挡期间受伤乘以 0.25；弓手蓄力满蓄时伤害乘以 2.15 且箭矢可穿透。"
        "敌人数值随层数缩放，Boss HP = 240 + stage×90。",
        indent=True,
    )

    add_title(doc, "第四章  功能模块设计与类结构", 1)
    add_body(doc, "系统按功能划分为以下模块，各模块列出主要类及其职责。", indent=True)
    modules = [
        ("4.1  启动与生命周期模块", "启动与生命周期模块主要类", [
            ("ModernRogueBootstrapper", "场景加载后创建 Inventory、Progression、Director 等；将 Player 从 3D 迁移到 2D"),
            ("GameBootstrapper", "Legacy 启动入口；Modern 模式下抑制旧 RunLoop"),
            ("ModernRogueLegacyUiSuppressor", "隐藏旧背包/角色面板，避免 UI 冲突"),
            ("GameStartMenuPanel", "主菜单：新游戏、继续、设置、帮助"),
        ]),
        ("4.2  地牢与关卡模块", "地牢与关卡模块主要类", [
            ("SoulKnightDirector", "核心导演：层数、HUD、Toast、结算、异步建关"),
            ("SoulKnightDungeonBuilder", "程序化生成 8–13 房间网格、走廊、门洞与大厅布局"),
            ("SoulKnightStageProfile", "五层主题、敌种概率、奖励权重、Boss 类型"),
            ("RoomTrigger2D", "进房锁门、多波刷怪、清房开宝箱/传送门"),
            ("DungeonCrawlCombatRoomVisual", "Dungeon Crawl 风格地板砖、墙、装饰（纯视觉）"),
            ("SoulKnightEnemyFactory", "生成史莱姆、精英史莱姆、骷髅弓手、四类 Boss"),
        ]),
        ("4.3  战斗与玩家模块", "战斗与玩家模块主要类", [
            ("PlayerController2D", "WASD 移动、闪避、瞄准方向"),
            ("PlayerAttack2D", "三职业攻击：战士扇形近战+格挡、弓手蓄力射击、法师法力弹"),
            ("PlayerClassVfx2D", "蓄力条、格挡、箭矢等职业 VFX"),
            ("ElementalCoreCombatSystem", "六元素虫核套装战斗效果"),
            ("LoopActors.cs", "敌人与 Boss AI"),
        ]),
        ("4.4  进度与 Meta 模块", "进度与 Meta 模块主要类", [
            ("RunProgressionSystem", "职业、宝物、升级三选一、护甲护盾、Boss 精华 Meta"),
            ("GameDataModel", "物品/虫核数据表、属性计算公式"),
            ("GameSaveService", "PlayerPrefs JSON 存档"),
            ("WeaponForgePanel", "大厅武器铺：消耗 Boss 精华升级"),
        ]),
        ("4.5  背包与 UI 模块", "背包与 UI 模块主要类", [
            ("NewInventorySystem", "24 格背包、装备槽、虫核、金币经验"),
            ("UIManager", "Tab 切换：背包/角色/虫核/宝物"),
            ("RuntimeUIVisuals", "程序化 UI 样式（边框、色块）"),
        ]),
        ("4.6  视觉与动画模块", "视觉与动画模块主要类", [
            ("Art2DUtility", "Resources 精灵加载、别名、透明边裁剪、Warm 预加载"),
            ("RuntimeSpriteAnimator2D", "代码驱动帧动画状态机"),
            ("SKCharacterVisual2D", "射手/法师 SK 风格视觉；弓手用 TinySwords"),
            ("TinySwordsPlayerVisual2D", "战士：TinySwords Prefab + Unity Animator"),
            ("TinySwordsExternalSpriteLibrary", "弓手 sprite sheet 解析（Resources/StreamingAssets）"),
            ("RogueActorVisual2D", "敌人/Boss 动画配置"),
        ]),
        ("4.7  音频模块", "音频模块主要类", [
            ("GameAudioService", "BGM 切换（Boss 战）、攻击/格挡/拉弓音效"),
        ]),
    ]
    for sec_title, cap, classes in modules:
        add_title(doc, sec_title, 2)
        ch = sec_title.split(".")[0]
        add_table(doc, ch, cap, ["类名", "职责"], classes)

    add_title(doc, "第五章  游戏规则与玩法说明", 1)
    add_title(doc, "5.1  基本规则", 2)
    add_bullet(doc, "游戏为单局 Roguelike：每次冒险从大厅出发，共 5 层地牢，每层含多个房间与 1 个 Boss 房。")
    add_bullet(doc, "进入房间深处触发战斗，四门落锁；清空所有敌人后解锁并出现奖励箱或传送门。")
    add_bullet(doc, "击杀敌人获得经验与金币；升级时三选一强化攻击/生命/防御/移速等。")
    add_bullet(doc, "可拾取虫核、宝物、消耗品；虫核装备后触发元素套装效果。")
    add_bullet(doc, "第 5 层 Boss 房为双 Boss 同场；通关后显示胜利结算。")
    add_title(doc, "5.2  职业与战斗", 2)
    add_table(
        doc, "5", "三职业战斗特性对比",
        ["职业", "武器特点", "特殊机制"],
        [
            ["战士", "近战扇形攻击（约 95°）", "格挡键 1 秒内减伤 75%；TinySwords 动画"],
            ["弓箭手", "远程蓄力射击", "蓄力 0–2 秒，满蓄 2.15× 伤害且穿透；闪避时长 ×1.35"],
            ["法师", "消耗法力发射魔法弹", "当前版本 UI 存在但选职未完全开放"],
        ],
    )
    add_title(doc, "5.3  操作说明", 2)
    add_table(
        doc, "5", "游戏操作说明",
        ["操作", "功能"],
        [
            ["W/A/S/D", "移动"],
            ["鼠标指向", "瞄准方向"],
            ["左键 / 攻击键", "普通攻击 / 蓄力（弓手）"],
            ["格挡键", "战士格挡"],
            ["闪避键", "翻滚闪避（弓手闪避时间更长）"],
            ["Tab / 背包键", "打开背包/角色/虫核/宝物面板"],
            ["M 键", "大厅武器铺（Meta 升级）"],
            ["E / 交互", "开宝箱、传送门、雕像选职等"],
        ],
    )
    add_title(doc, "5.4  如何开始游戏", 2)
    add_body(
        doc,
        "开发者在 Unity 中打开 Main.unity 场景，点击 Play 即可在 Editor 内试玩。"
        "对外分发时使用菜单 Modern Rogue → Build Windows Release 生成 Builds/Windows 文件夹，"
        "将整文件夹压缩后发给他人；对方解压后双击「黑暗地牢.exe」即可运行（需 Windows 64 位）。",
        indent=True,
    )

    add_title(doc, "第六章  游戏难度、构筑与进度系统", 1)

    add_title(doc, "6.1  局外 Meta 构筑与局内 Run 构筑", 2)
    add_body(
        doc,
        "《黑暗地牢》采用「局外永久成长 + 局内单局成长」的双层构筑体系："
        "局外 Meta 构筑在多次 Run 之间累积，改变后续冒险的起点强度；"
        "局内 Run 构筑仅在当前冒险有效，死亡或通关后大部分局内强化清零，"
        "但 Boss 精华等 Meta 资源会保留。二者相互独立又形成长期目标。",
        indent=True,
    )
    add_table(
        doc, "6", "局外 Meta 构筑与局内 Run 构筑对比",
        ["维度", "局外 Meta 构筑", "局内 Run 构筑"],
        [
            ["生效范围", "跨 Run 永久保留", "仅当前冒险有效"],
            ["主要货币", "Boss 精华（击败 Boss 获得）", "金币、经验、房间掉落"],
            ["核心界面", "大厅武器铺（M 键，WeaponForgePanel）", "升级三选一、Tab 面板、局内商人"],
            ["主要收益", "武器/防具等级提升；每 5 级 Meta 随机词条；防具 Lv.5 护甲护盾", "攻击/生命/防御/移速；宝物；虫核；消耗品"],
            ["职业与装备", "Meta 武器/防具属性叠加到每局起点", "每局固定职业 loadout，局内不可换主武器"],
            ["相关类", "WeaponForgePanel、RunProgressionSystem、GameSaveService", "NewInventorySystem、RunProgressionSystem、RogueShopPanel"],
        ],
    )
    add_title(doc, "6.1.1  局外 Meta 构筑（武器铺）", 3)
    add_body(
        doc,
        "玩家在大厅按 M 键打开武器铺（图 6-1），消耗 Boss 精华升级武器或防具。"
        "每次升级消耗等于「当前等级 + 1」颗精华；武器每升一级固定 +1 攻击，"
        "防具每升一级 +1 防御并 +3 最大生命。"
        "当武器或防具达到 5 级、10 级等 5 的倍数时，会随机获得一条 Meta 词条（如吸血、暴击、攻速、减伤等）；"
        "防具 Meta 达到 5 级时额外解锁「护甲护盾」——独立于 HP 的伤害吸收层，"
        "被击破后延迟回复。Meta 等级与词条通过 GameSaveService 写入 PlayerPrefs，"
        "下次「新游戏」或「继续游戏」时自动加载。",
        indent=True,
    )
    add_figure(doc, "6", "大厅武器铺 Meta 升级界面", "武器铺界面.png", 14)

    add_title(doc, "6.1.2  局内 Run 构筑", 3)
    add_body(
        doc,
        "进入地牢后的局内构筑来源包括："
        "（1）**升级三选一**：经验达到 28+level×14 时升级，随机提供攻/血/防/移速等选项，每 5 级追加稀有项；"
        "（2）**宝物**：清房或特殊房获得，提供元素碎片、商店刷新、复活、无伤成长等被动；"
        "（3）**虫核**：拾取并装备至虫核槽（见 6.3 节），同元素 2/3/4 件套触发战斗特效；"
        "（4）**局内商人**：部分房间标记为 Shop，玩家用本局金币购买消耗品、虫核或装备（图 6-2）；"
        "（5）**房间奖励箱**：按 RoomRewardType 掉落金币、生命、虫核或宝物。",
        indent=True,
    )
    add_figure(doc, "6", "局内商人购买界面（我的物品 / 商人货物）", "商人界面.png", 14)

    add_title(doc, "6.2  五层难度曲线", 2)
    add_table(
        doc, "6", "五层地牢难度配置",
        ["层数", "主题", "房间数", "波次", "敌人特点", "Boss"],
        [
            ["1", "废墟前厅", "8–11", "3", "远程 18%，精英 2%", "地牢巨像 Titan"],
            ["2", "箭廊墓道", "9–12", "3", "远程 72%", "风暴守卫 StormGuard"],
            ["3", "腐潮深井", "9–12", "3", "精英 24%", "巢穴母虫 BroodQueen"],
            ["4", "熔炉回廊", "10–13", "4", "高宝物/装备权重", "烬火术士 EmberMage"],
            ["5", "王座外环", "8–10", "4", "敌数 5–7", "双 Boss 随机组合"],
        ],
    )
    add_title(doc, "6.3  局内经验与 Meta 精华", 2)
    add_body(
        doc,
        "局内经验需求公式：28 + level × 14；升级时从 RunProgressionSystem 提供的候选中三选一。"
        "Meta 精华仅由击败 Boss 掉落，与局内金币相互独立——"
        "前者驱动武器铺跨局成长，后者驱动单局商人消费与角色页少量属性购买。",
        indent=True,
    )
    add_title(doc, "6.4  虫核元素套装", 2)
    add_body(
        doc,
        "虫核按元素分为风、火、水、雷、金五种（土之碎片仅用于宝物解锁额外镶嵌位，无战斗套装效果）。"
        "玩家最多装备 6 个虫核（同元素碎片宝物可额外 +2 槽），"
        "同一元素装备 2/3/4 颗时分别触发下表套装效果，由 ElementalCoreCombatSystem 在战斗事件中检测并执行。",
        indent=True,
    )
    add_table(
        doc, "6", "虫核元素套装效果（依据 ElementalCoreCombatSystem 实现）",
        ["元素", "2 件套", "3 件套", "4 件套"],
        [
            ["风", "移速 +10%；伤害随移速加成提升", "移速 +20%；闪避后下次攻击必暴击", "移速 +30%；清房获得免费闪避充能（约 3.2s 冷却）"],
            ["火", "攻击 45% 概率生成环绕火球", "火球可发射追踪；触发概率 60%", "触发概率 80%；火球伤害提升且附带燃烧"],
            ["水", "移动时留下水渍", "站在水渍上持续回血；敌人踩水渍获得湿润减速", "对湿润敌人伤害 +8%"],
            ["雷", "攻击 50% 概率施加感电", "感电概率 80%；链电伤害相邻敌人；链电后 2s 攻速 +15%", "单目标亦可蓄力释放；感电目标受伤害逐层加深"],
            ["金", "击杀 +1 金币；清房 +10 金币（完美 +15）", "击杀 +2 金币；商店九折；处决残血（≤4% HP）敌人", "25% 概率触发落金打击额外伤害"],
        ],
    )

    add_title(doc, "第七章  图形渲染、场景与动画实现", 1)

    add_title(doc, "7.1  2D 渲染管线与安全大厅", 2)

    add_title(doc, "7.1.1  2D 渲染与场景组织", 3)
    add_body(
        doc,
        "本游戏采用纯 2D 渲染管线：场景中可见对象统一使用 SpriteRenderer 组件显示，"
        "主相机为正交投影俯视角，通过 Sorting Layer 与 sortingOrder 控制玩家、敌人、"
        "环境装饰与 UI 的绘制先后顺序。",
        indent=True,
    )
    add_body(
        doc,
        "在场景组织上，项目遵循「碰撞逻辑与视觉表现分离」原则："
        "战斗房间中，DungeonCrawlCombatRoomVisual 负责按 1×1 网格逐格铺设地板与墙体精灵，"
        "而 BoxCollider2D 碰撞壳由 SoulKnightDungeonBuilder 独立生成，便于在不影响玩法的前提下迭代美术。"
        "大厅与战斗房均不在编辑器中手工逐房摆放，而是由运行时脚本根据布局数据动态实例化视觉节点。",
        indent=True,
    )

    add_title(doc, "7.1.2  安全大厅的功能定位与可玩循环", 3)
    add_body(
        doc,
        "安全大厅（Lobby，图 7-1）是玩家每次 Run 的起点，也是局外 Meta 与局内冒险的衔接枢纽。"
        "与《元气骑士》等作品中的 Hub 类似，大厅在单屏空间内集中呈现本游戏的核心交互闭环，"
        "使玩家在进 dungeon 前即可完成职业配置、资源准备与长期成长决策。",
        indent=True,
    )
    add_body(
        doc,
        "大厅可玩循环按推荐路径可概括为："
        "（1）开启初始宝箱获取基础补给；"
        "（2）在职业雕像处选择或切换战士/弓箭手；"
        "（3）在武器铺消耗 Boss 精华升级武器与防具（详见第六章 6.1.1）；"
        "（4）经传送门进入当前层程序化地牢。"
        "其中，职业雕像与武器铺分别由 LobbyInteractables、WeaponForgePanel 驱动，"
        "传送门则与 SoulKnightDirector 的层进度状态联动。",
        indent=True,
    )
    add_table(
        doc, "7", "安全大厅主要交互要素",
        ["区域/对象", "交互方式", "功能说明", "相关类"],
        [
            ["初始宝箱", "E 键", "首次进入提供基础物品与补给", "SoulKnightChest / LobbyInteractables"],
            ["职业雕像", "E 键", "战士/弓箭手转职，确认本局职业", "LobbyInteractables / RunProgressionSystem"],
            ["武器铺", "M 键", "消耗 Boss 精华进行 Meta 武器/防具升级", "WeaponForgePanel"],
            ["地牢传送门", "E 键", "进入当前层程序化生成的地牢地图", "SoulKnightPortal / SoulKnightDirector"],
            ["训练靶（可选）", "靠近", "大厅内测试攻击与手感", "LobbyInteractables"],
        ],
    )
    add_body(
        doc,
        "从图形学角度看，大厅同时承担「环境展示」与「交互引导」两类任务："
        "一方面需通过地砖、立柱、雕像等精灵营造可识别的空间语义；"
        "另一方面需保证各交互点在俯视角下具有足够对比度与可读性，"
        "避免玩家误判可行走区域与装饰边界。",
        indent=True,
    )

    add_title(doc, "7.1.3  大厅视觉实现", 3)
    add_body(
        doc,
        "大厅视觉由 DungeonCrawlBaseRoomRuntimeVisual 在运行时挂载："
        "地面与墙体砖块引用 ArtIntegration/Environment/Lobby 下的精灵资源，"
        "职业雕像、传送门、武器铺等交互物则按预设锚点放置，"
        "并与碰撞体、触发器对齐以保证交互位置准确。"
        "图 7-1 展示了大厅全景，可见 Dungeon Crawl 风格地砖、"
        "中央三职业石像、两侧功能区域及通往地牢的传送门，"
        "对应上文所述的可玩循环各节点。",
        indent=True,
    )
    add_figure(
        doc,
        "7",
        "安全大厅全景：地砖铺设、职业雕像、武器铺与地牢传送门",
        "7-1大厅.png",
        14,
    )
    add_title(doc, "7.2  战斗场景与人物动画", 2)
    add_body(
        doc,
        "进入地牢后，战斗房间同样由 DungeonCrawlCombatRoomVisual 程序化铺砖，"
        "与大厅共享 Dungeon Crawl 美术风格，但在色调与装饰密度上随 StageProfile 按层变化。",
        indent=True,
    )
    add_body(
        doc,
        "角色动画采用双轨方案：战士由 TinySwordsPlayerVisual2D 实例化 Prefab，"
        "通过 Unity Animator 状态机驱动 Idle/Run/Attack/Guard；"
        "弓箭手由 SKCharacterVisual2D 配合 RuntimeSpriteAnimator2D，"
        "对 TinySwords 弓手 sprite sheet 动态切片换帧；"
        "敌人与 Boss 则由 RogueActorVisual2D 加载 Art2D/AnimationSprites 序列帧，"
        "Boss 战另叠加火圈、闪电等程序化 VFX。",
        indent=True,
    )
    add_figure(doc, "7", "TinySwords 战士角色运行时显示", "7-2运行时战士显示.png", 14)
    add_body(
        doc,
        "图 7-3 由三张运行截图拼接而成，分别展示箭矢投射物、弓箭手蓄力姿态与 Boss 战中的技能特效，"
        "用于验证精灵透明通道抠图质量及在游戏场景中的实际显示效果。",
        indent=True,
    )
    add_figure(doc, "7", "透明通道抠图与精灵质量验证（箭矢、蓄力、Boss 特效）", "fig7-3_composite.png", 14)
    add_title(doc, "7.3  UI 与资源发布", 2)
    add_body(
        doc,
        "Modern UI 由 RuntimeUIVisuals、UIManager 等在运行时动态创建，采用 Tab 切换背包、角色、虫核与宝物四个面板。"
        "打开面板时通过 ModalPauseToken 暂停游戏逻辑，并以半透明遮罩覆盖场景。"
        "图 7-4 为四张 UI 截图拼接，展示各子面板的布局与信息层次。",
        indent=True,
    )
    add_body(
        doc,
        "弓手图集通过 StreamingAssets + Resources 双路径加载；"
        "ReleaseBuildPreparer 在构建前同步关键资源，保证 Editor 与 Release 表现一致。",
        indent=True,
    )
    add_figure(
        doc,
        "7",
        "Modern UI 四 Tab 面板（背包、角色、虫核、宝物）",
        "fig7-4_composite.png",
        14,
    )

    add_title(doc, "第八章  创新点与同类游戏对比", 1)
    add_table(
        doc, "8", "与同类 Roguelike 游戏对比",
        ["维度", "典型 Roguelike（如元气骑士）", "本项目《黑暗地牢》"],
        [
            ["房间结构", "固定模板随机连接", "StageProfile 驱动五层主题+签名房间"],
            ["成长", "武器/天赋为主", "局内升级+宝物+虫核套装+Meta 武器铺四层叠加"],
            ["Boss", "单层单 Boss", "第 5 层双 Boss 同场"],
            ["背包", "简化或武器切换", "24 格背包+虫核槽+装备 Meta 升级"],
            ["视觉", "统一像素风", "Dungeon Crawl 砖块+TinySwords 混合集成"],
            ["存档", "整局存档为主", "PlayerPrefs 细粒度 Run 存档"],
        ],
    )
    add_title(doc, "8.1  创新玩法总结", 2)
    add_bullet(doc, "六元素虫核 2/3/4 件套：风、火、水、雷、金不同机制联动。")
    add_bullet(doc, "完美清房机制：进房后不掉血清房触发额外奖励。")
    add_bullet(doc, "Meta 护甲护盾：独立于 HP 的吸收层。")
    add_bullet(doc, "双动画管线：Mecanim 与代码帧动画按职业分流。")

    add_title(doc, "第九章  运行效果与测试", 1)
    add_body(
        doc,
        "本章展示游戏从主菜单、职业选择到俯视角战斗的典型运行界面。"
        "安全大厅见第七章 7.1；程序化地牢拓扑见第三章 3.4；"
        "局外武器铺与局内商人见第六章 6.1；"
        "胜负结算见 9.5 节。"
        "多层 Boss 战与正式包实机演示留待答辩现场展示。",
        indent=True,
    )

    add_title(doc, "9.1  主菜单与各选项说明", 2)
    add_body(
        doc,
        "游戏启动后进入 GameStartMenuPanel 主菜单（图 9-1）。"
        "主菜单由 RuntimeUIVisuals 动态构建，提供四项入口，"
        "各选项功能如下表所示。",
        indent=True,
    )
    add_table(
        doc, "9", "主菜单选项说明",
        ["选项", "功能说明", "备注"],
        [
            ["新游戏", "清空当前 Run 进度，进入安全大厅开始全新冒险", "会重置局内等级/背包，保留 Meta 武器铺等级与 Boss 精华"],
            ["继续上次游戏", "读取 GameSaveService 存档，恢复层数、位置、背包与虫核等", "无存档时按钮灰色不可点"],
            ["游戏说明", "打开帮助面板，含基础操作、五层主题、Boss 介绍、面板与存档说明", "内置多 Tab 侧边栏"],
            ["设置", "调整攻击/闪避/背包等按键绑定及 BGM 音量", "支持恢复默认键位"],
        ],
    )
    add_figure(doc, "9", "游戏开始界面（主菜单）", "9-1游戏开始界面.png", 14)

    add_title(doc, "9.2  职业选择", 2)
    add_body(
        doc,
        "在大厅与职业雕像交互后弹出职业选择界面（图 9-2）。"
        "当前版本开放战士与弓箭手；选择结果写入 RunProgressionSystem.CurrentClass，"
        "并决定本局攻击方式、视觉方案（TinySwords 战士 / 弓手帧动画）及初始 loadout。"
        "职业可在同一 Run 内于雕像处再次切换。",
        indent=True,
    )
    add_figure(doc, "9", "职业选择界面（战士 / 弓箭手）", "9-2职业选择.png", 14)

    add_title(doc, "9.3  俯视角战斗场景", 2)
    add_body(
        doc,
        "图 9-3 为地牢战斗房内的 2D 俯视角实机画面，"
        "可见 Dungeon Crawl 风格地砖、玩家与敌人、投射物及 HUD 信息，"
        "体现第三章所述程序化房间与第七章所述 2D 渲染管线在实际游玩中的效果。",
        indent=True,
    )
    add_figure(doc, "9", "2D 俯视角战斗场景", "9-3俯视角战斗.png", 14)

    add_title(doc, "9.5  冒险结算界面（通关成功 / 冒险失败）", 2)
    add_body(
        doc,
        "单局 Run 结束时，SoulKnightDirector 调用 ShowRunResultPanel 弹出结算面板，"
        "并通过 ModalPauseToken 将 Time.timeScale 置 0，暂停后台战斗逻辑。"
        "结算 UI 同样由 RuntimeUIVisuals 运行时构建，"
        "根据胜负切换标题色、边框色与统计项排列顺序。",
        indent=True,
    )
    add_body(
        doc,
        "**通关成功**（图 9-4）：玩家在第 5 层 Boss 房清敌后，"
        "经 Victory 传送门触发 ShowSettlement。"
        "面板标题为「通关成功！」，副标题提示魔法石已夺回；"
        "统计区展示获得金币、击杀怪物、角色等级与通关层数（1-5），"
        "点击「返回大厅」回到安全 Hub，本局局内成长清零，"
        "但 Boss 精华及武器铺 Meta 等级等局外构筑数据仍保留在存档中。",
        indent=True,
    )
    add_figure(doc, "9", "通关成功结算界面", "通关成功.png", 14)
    add_body(
        doc,
        "**冒险失败**（图 9-5）：玩家 HP 归零且复活宝物/消耗品均不可用后，"
        "PlayerLifeState2D 调用 ShowDeathSettlement。"
        "面板标题为「冒险失败」，副标题说明局外收获仍会保留；"
        "统计区展示到达层数、击杀数、角色等级与本局金币；"
        "点击「返回基地」回到大厅，"
        "RunProgressionSystem.EndRun 会结束当前冒险并清除局内进度，"
        "已获得的 Boss 精华不会丢失，可在下次 Run 前继续于武器铺消费。",
        indent=True,
    )
    add_figure(doc, "9", "冒险失败结算界面", "通关失败.png", 14)

    add_title(doc, "9.4  测试与发布", 2)
    add_body(
        doc,
        "测试内容包括：五层地牢连通性、Boss 战、胜负结算面板、存档续玩、战士战斗与 UI 交互、"
        "Windows Release 包资源加载一致性等。"
        "发布流程由 ReleaseBuildPreparer 自动化构建 StandaloneWindows64 可执行文件；"
        "部分动态流程（多层地牢切换、Boss 战过程）留待答辩现场演示。",
        indent=True,
    )

    add_title(doc, "第十章  结论与展望", 1)
    add_title(doc, "10.1  结论", 2)
    add_body(
        doc,
        "本文完成了基于 Unity 的 2D Roguelike 游戏《黑暗地牢》的设计与实现。"
        "项目在图形学层面实践了精灵渲染、序列帧动画、透明贴图处理、"
        "运行时资源加载与双路径发布兼容；在系统设计上实现了模块化架构、"
        "程序化地牢、多层进度与虫核元素战斗。游戏已具备完整的可玩闭环与 Windows 独立发布能力，"
        "达到了计算机图形学大作业对「设计—实现—展示」的综合要求。",
        indent=True,
    )
    add_title(doc, "10.2  不足与展望", 2)
    add_bullet(doc, "法师职业尚未完全开放，后续可补齐法力系统与专属 VFX。")
    add_bullet(doc, "战斗房间视觉可继续向 Dungeon Crawl 风格统一。")
    add_bullet(doc, "可增加更多 Boss 机制与联机观战/排行榜等扩展。")
    add_bullet(doc, "图形方面可探索 Shader 受击闪白、Bloom 与后处理以增强打击感。")

    add_title(doc, "参考文献", 1)
    for r in [
        "[1] Unity Technologies. Unity Manual: 2D Sprite [EB/OL]. https://docs.unity3d.com/Manual/Sprites.html",
        "[2] Unity Technologies. Unity Manual: Physics 2D [EB/OL]. https://docs.unity3d.com/Manual/Physics2D.html",
        "[3] 冯乐乐. Unity 游戏开发实战[M]. 人民邮电出版社.",
        "[4] Togelius J, et al. Procedural Content Generation in Games[M]. Springer, 2016.",
        "[5] 项目源码：Assets/Scripts/ModernRogue/ 及各 Stage 4.x 集成文档.",
    ]:
        add_body(doc, r, indent=False)

    doc.add_page_break()
    add_title(doc, "附录 A  ModernRogue 主要类清单", 1)
    add_body(doc, "以下为 Assets/Scripts/ModernRogue/ 目录下主要类名汇总。", indent=True)
    class_list = [
        "SoulKnightDirector", "SoulKnightDungeonBuilder", "SoulKnightStageProfile",
        "RoomTrigger2D", "SoulKnightEnemyFactory", "DungeonCrawlCombatRoomVisual",
        "ModernRogueBootstrapper", "GameStartMenuPanel", "RunProgressionSystem",
        "PlayerController2D", "PlayerAttack2D", "ElementalCoreCombatSystem", "LoopActors",
        "NewInventorySystem", "UIManager", "Art2DUtility", "RuntimeSpriteAnimator2D",
        "SKCharacterVisual2D", "TinySwordsPlayerVisual2D", "TinySwordsExternalSpriteLibrary",
        "RogueActorVisual2D", "GameAudioService", "GameSaveService",
    ]
    add_body(doc, "、".join(class_list), indent=False)

    legacy_copy = os.path.join(REPORT_ROOT, "计算机图形学大作业报告_黑暗地牢_新版.docx")
    if os.path.isfile(legacy_copy):
        os.remove(legacy_copy)
        print("Removed legacy copy:", legacy_copy)

    temp_out = OUTPUT + ".tmp.docx"
    doc.save(temp_out)
    try:
        if os.path.isfile(OUTPUT):
            os.remove(OUTPUT)
        os.replace(temp_out, OUTPUT)
        out = OUTPUT
    except PermissionError:
        if os.path.isfile(OUTPUT):
            os.remove(temp_out)
            raise SystemExit(
                "无法覆盖报告：请先关闭 Word 中的「计算机图形学大作业报告_黑暗地牢.docx」后重试。"
            )
        os.replace(temp_out, OUTPUT)
        out = OUTPUT
    print("Saved:", out)
    print("Tables:", sum(_table_idx.values()), "Figures:", sum(_figure_idx.values()))


if __name__ == "__main__":
    os.makedirs(REPORT_ROOT, exist_ok=True)
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    sync_screenshots_to_report_folder()
    build_figure_7_3_composite()
    build_figure_7_4_composite()
    build_report()
